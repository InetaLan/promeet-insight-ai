from pathlib import Path
from docx import Document


def read_docx(path: str | Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_transcript_file(path: str | Path) -> str:
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported file type: {path.suffix}")
